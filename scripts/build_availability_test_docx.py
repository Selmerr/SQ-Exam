import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "software-quality" / "04-test-design" / "availability-unit-test-design.md"
DEFAULT_OUTPUT = ROOT / "software-quality" / "04-test-design" / "availability-unit-test-design.docx"
OUTPUT = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "8A8A8A")


def style_table(table):
    set_table_borders(table)
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 78, 121)


def add_code_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)


def add_paragraph_with_code(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    parts = text.split("`")
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        if idx % 2 == 1:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 45, 130)
        else:
            run.font.size = Pt(10)


def add_table(document, rows):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            parts = value.split("`")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.font.name = "Consolas"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                    run.font.size = Pt(8.3)
                    run.font.color.rgb = RGBColor(100, 45, 130)
                else:
                    run.font.size = Pt(8.5)
    style_table(table)
    document.add_paragraph()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(set(cell) <= {"-", ":", " "} for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build_docx():
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    for style_name, size, color in [
        ("Title", 20, RGBColor(31, 78, 121)),
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 12, RGBColor(70, 70, 70)),
        ("Heading 3", 10.5, RGBColor(90, 90, 90)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    title = document.add_paragraph()
    title.style = styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Availability Unit Test Design Based on Black-Box Techniques")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.add_run("Software Quality - Unit test design and traceability").italic = True

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    i = 1  # skip duplicate markdown title
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if line.startswith("```"):
            if in_code:
                for code_line in code_lines:
                    if code_line.strip():
                        add_code_paragraph(document, code_line)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line:
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                add_table(document, rows)
            continue

        if line.startswith("# "):
            p = document.add_paragraph(line[2:], style="Heading 1")
        elif line.startswith("## "):
            p = document.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            p = document.add_paragraph(line[4:], style="Heading 2")
        else:
            add_paragraph_with_code(document, line)
        i += 1

    document.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
